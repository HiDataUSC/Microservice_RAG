from services.common.AWS_handler import S3Handler

import os
import json
import uuid
from uuid import uuid4
from abc import ABC, abstractmethod
from docx import Document as DocxReader
from pathlib import Path
import mimetypes
import zipfile
from langchain import hub
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_chroma import Chroma
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough
from langchain_core.documents import Document
from langchain_core.prompts import ChatPromptTemplate
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain.storage import InMemoryByteStore
from langchain.retrievers.multi_vector import MultiVectorRetriever


# Abstract base class defining methods for file processing states
class FileProcessingState(ABC):
    # Abstract method to read file content
    @abstractmethod
    def read(self, file_path):
        """Read the file."""
        pass
    
    # Abstract method for content preprocessing
    @abstractmethod
    def preprocess(self, content):
        """Preprocess the content."""
        pass

    # Abstract method for vectorizing content
    @abstractmethod
    def vectorize(self, content, local_folder):
        """Vectorize the preprocessed content."""
        pass
    
    # Abstract method for local storage of vectorized data
    @abstractmethod
    def store_local(self, vectorstore, local_folder):
        """Store the vectorized data locally."""
        pass
    
    # Abstract method for cloud storage of data
    @abstractmethod
    def store_cloud(self):
        """Store data to the cloud"""
        pass

# State for processing text (.txt) files
class TextFileState(FileProcessingState):
    def __init__(self):
        self.embeddings = OpenAIEmbeddings()  # Initialize OpenAI embeddings
        self.file_path = None
        self.vectorstore = None
        self.local_folder = None
        self.doc_id = None

    # Read method to read content from a text file
    def read(self, file_path):
        """Read the text file."""
        self.file_path = file_path
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()

    # Preprocess content by summarizing it
    def preprocess(self, content):
        doc = Document(page_content=content)
        chain = (
            {"doc": lambda x: x.page_content}
            | ChatPromptTemplate.from_template("Summarize the following document:\n\n{doc}")
            | ChatOpenAI(model="gpt-3.5-turbo", max_retries=0)
            | StrOutputParser()
        )
        summary = chain.invoke(doc)  # Generate summary using LLM
        return summary

    # Vectorize the text content and set up local folder for persistence
    def vectorize(self, local_folder):
        """Vectorize the text chunks."""
        self.local_folder = local_folder
        self.vectorstore = Chroma(
            collection_name="summaries", 
            embedding_function=self.embeddings,
            persist_directory=self.local_folder
        )
        return self.vectorstore
    
    # Store vectorized document locally with metadata
    def store_local(self, doc_id, content):
        self.doc_id = doc_id
        summary_docs = [
            Document(page_content=content, metadata={"doc_id": doc_id, "doc_type": "txt"})
        ]
        
        try:
            self.vectorstore.add_documents(summary_docs, ids=[doc_id])  # Add documents to vector store
            # Save document IDs locally to a JSON file for tracking
            doc_id_file_path = os.path.join(self.local_folder, "doc_id.json")
            if not os.path.exists(doc_id_file_path):
                with open(doc_id_file_path, 'w', encoding='utf-8') as f:
                    json.dump([], f)
            with open(doc_id_file_path, 'r', encoding='utf-8') as f:
                doc_ids = json.load(f)
            if doc_id not in doc_ids:
                doc_ids.append(doc_id)
            with open(doc_id_file_path, 'w', encoding='utf-8') as f:
                json.dump(doc_ids, f, indent=4)
        except Exception as e:
            print(f"Error occurred in file_processing_stats.store_local: {str(e)}")

    # Upload the original file to cloud storage
    def store_cloud(self):
        """Upload the original Word file to cloud storage."""
        s3_handler = S3Handler()
        ext = Path(self.file_path).suffix.lower()
        object_name = f"{self.doc_id}{ext}"
        file_name = Path(self.file_path).name
        s3_handler.upload_file(self.file_path, folder_prefix="files", object_name=object_name, metadata={"name": file_name})

        useful_files = ["chroma.sqlite3", "doc_id.json"]
        for root, _, files in os.walk(self.local_folder):
            for file in files:
                if file in useful_files:
                    file_path = os.path.join(root, file)
                    s3_handler.upload_file(file_path, folder_prefix="vectorized_db", object_name=file)



# Placeholder class for PDF file processing
class PDFFileState(FileProcessingState):
    def read(self, file_path):
        """Read the file."""
        pass
    
    def preprocess(self, content):
        """Preprocess the content."""
        pass

    def vectorize(self, content, local_folder):
        """Vectorize the preprocessed content."""
        pass
    
    def store_local(self, vectorstore, local_folder):
        """Store the vectorized data locally."""
        pass
    
    def retrieve(self, query, local_folder):
        """Retrieve using similarity search."""
        pass


# State for processing Word (.docx) files
class WordFileState(FileProcessingState):
    def __init__(self):
        self.embeddings = OpenAIEmbeddings()  # Initialize embeddings
        self.file_path = None
        self.vectorstore = None
        self.local_folder = None
        self.doc_id = None

    # Read content from a Word document using python-docx
    def read(self, file_path):
        """Read the Word file."""
        self.file_path = file_path
        doc = DocxReader(file_path)
        full_text = [para.text for para in doc.paragraphs]
        return '\n'.join(full_text)

    # Preprocess content by summarizing it using LLM
    def preprocess(self, content):
        doc = Document(page_content=content)
        chain = (
            {"doc": lambda x: x.page_content}
            | ChatPromptTemplate.from_template("Summarize the following document:\n\n{doc}")
            | ChatOpenAI(model="gpt-3.5-turbo", max_retries=0)
            | StrOutputParser()
        )
        summary = chain.invoke(doc)
        return summary

    # Vectorize content and set up local folder for persistence
    def vectorize(self, local_folder):
        """Set up the vector store for embeddings."""
        self.local_folder = local_folder
        self.vectorstore = Chroma(
            collection_name="summaries",
            embedding_function=self.embeddings,
            persist_directory=self.local_folder
        )
        return self.vectorstore

    # Store vectorized document locally with metadata
    def store_local(self, doc_id, content):
        """Store the document vectors locally with metadata."""
        self.doc_id = doc_id
        summary_docs = [
            Document(page_content=content, metadata={"doc_id": doc_id, "doc_type": "docx"})
        ]
        try:
            self.vectorstore.add_documents(summary_docs, ids=[doc_id])  # Add to vector store
            doc_id_file_path = os.path.join(self.local_folder, "doc_id.json")
            if not os.path.exists(doc_id_file_path):
                with open(doc_id_file_path, 'w', encoding='utf-8') as f:
                    json.dump([], f)
            with open(doc_id_file_path, 'r', encoding='utf-8') as f:
                doc_ids = json.load(f)
            if doc_id not in doc_ids:
                doc_ids.append(doc_id)
            with open(doc_id_file_path, 'w', encoding='utf-8') as f:
                json.dump(doc_ids, f, indent=4)
        except Exception as e:
            print(f"Error occurred in WordFileState.store_local: {str(e)}")

    # Upload the original Word file to cloud storage
    def store_cloud(self):
        """Upload the original Word file to cloud storage."""
        s3_handler = S3Handler()
        ext = Path(self.file_path).suffix.lower()
        object_name = f"{self.doc_id}{ext}"
        file_name = Path(self.file_path).name
        s3_handler.upload_file(self.file_path, folder_prefix="files", object_name=object_name, metadata={"name": file_name})

        useful_files = ["chroma.sqlite3", "doc_id.json"]
        for root, _, files in os.walk(self.local_folder):
            for file in files:
                if file in useful_files:
                    file_path = os.path.join(root, file)
                    s3_handler.upload_file(file_path, folder_prefix="vectorized_db", object_name=file)


# Utility function to detect file type and return the appropriate state class
def detect_file_type(file_path):
    """Detects file type based on extension."""
    ext = Path(file_path).suffix.lower()
    mime_type, _ = mimetypes.guess_type(file_path)

    if ext == '.txt' or mime_type == 'text/plain':
        return TextFileState()
    elif ext == '.pdf' or mime_type == 'application/pdf':
        return PDFFileState()
    elif ext == '.docx' or mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
        return WordFileState()
    else:
        return None
