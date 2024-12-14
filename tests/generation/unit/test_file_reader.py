import os
import tempfile
import pytest
from docx import Document
from services.Text_Generation.file_reader import FileReader

@pytest.fixture
def setup_test_files(tmpdir):
    """Create temporary .txt and .docx files for testing"""
    # Create a .txt file
    txt_file = tmpdir.join("test_file.txt")
    txt_file.write("This is a test text file.")

    # Create a .docx file
    docx_file_path = tmpdir.join("test_file.docx")
    doc = Document()
    doc.add_paragraph("This is a test DOCX file.")
    doc.save(docx_file_path)

    return tmpdir

def test_read_txt_file(setup_test_files):
    """Test if the .txt file is read correctly"""
    directory_path = str(setup_test_files)
    context = FileReader.read_files(directory_path)

    assert "This is a test text file." in context

def test_read_docx_file(setup_test_files):
    """Test if the .docx file is read correctly"""
    directory_path = str(setup_test_files)
    context = FileReader.read_files(directory_path)

    assert "This is a test DOCX file." in context

def test_empty_directory(tmpdir):
    """Test reading from an empty directory"""
    directory_path = str(tmpdir)
    context = FileReader.read_files(directory_path)

    assert context == ""

def test_mixed_files(setup_test_files, tmpdir):
    """Test a directory with both valid and invalid file types"""
    # Create an unsupported .pdf file
    pdf_file = tmpdir.join("test_file.pdf")
    pdf_file.write("This is a PDF file, unsupported format.")

    # Reuse the setup_test_files
    directory_path = str(setup_test_files)
    context = FileReader.read_files(directory_path)

    assert "This is a test text file." in context
    assert "This is a test DOCX file." in context
    assert "This is a PDF file, unsupported format." not in context

def test_large_txt_file(tmpdir):
    """Test reading a large .txt file"""
    large_txt_file = tmpdir.join("large_test_file.txt")
    large_content = "A" * 10000  # 10,000 characters
    large_txt_file.write(large_content)

    directory_path = str(tmpdir)
    context = FileReader.read_files(directory_path)

    assert large_content in context

def test_empty_txt_file(tmpdir):
    """Test reading an empty .txt file"""
    empty_txt_file = tmpdir.join("empty_test_file.txt")
    empty_txt_file.write("")

    directory_path = str(tmpdir)
    context = FileReader.read_files(directory_path)

    assert context == ""

def test_empty_docx_file(tmpdir):
    """Test reading an empty .docx file"""
    empty_docx_file = tmpdir.join("empty_test_file.docx")
    doc = Document()
    doc.save(empty_docx_file)

    directory_path = str(tmpdir)
    context = FileReader.read_files(directory_path)

    assert context == ""

def test_multiple_txt_files(tmpdir):
    """Test reading multiple .txt files in the same directory"""
    txt_file_1 = tmpdir.join("test_file_1.txt")
    txt_file_1.write("This is the first text file.")
    txt_file_2 = tmpdir.join("test_file_2.txt")
    txt_file_2.write("This is the second text file.")

    directory_path = str(tmpdir)
    context = FileReader.read_files(directory_path)

    assert "This is the first text file." in context
    assert "This is the second text file." in context

def test_multiple_docx_files(tmpdir):
    """Test reading multiple .docx files in the same directory"""
    docx_file_1 = tmpdir.join("test_file_1.docx")
    doc1 = Document()
    doc1.add_paragraph("First DOCX file content.")
    doc1.save(docx_file_1)

    docx_file_2 = tmpdir.join("test_file_2.docx")
    doc2 = Document()
    doc2.add_paragraph("Second DOCX file content.")
    doc2.save(docx_file_2)

    directory_path = str(tmpdir)
    context = FileReader.read_files(directory_path)

    assert "First DOCX file content." in context
    assert "Second DOCX file content." in context

def test_files_with_special_characters(tmpdir):
    """Test reading files with special characters in their content"""
    special_txt_file = tmpdir.join("special_chars_test_file.txt")
    special_txt_file.write("Text with special characters: !@#$%^&*()")

    directory_path = str(tmpdir)
    context = FileReader.read_files(directory_path)

    assert "Text with special characters: !@#$%^&*()" in context
